import os
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, session
from werkzeug.security import generate_password_hash, check_password_hash
from models import db, User, Informe
from docx import Document
from docx.shared import Inches
from flask_sqlalchemy import SQLAlchemy
import secrets

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), "uploads")
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app = Flask(__name__)
app.secret_key = 'super-secreto'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///app_informe.db'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
db.init_app(app)

# Inicializa la base de datos la primera vez
with app.app_context():
    db.create_all()
    # Crea el usuario admin por defecto si no existe
    if not User.query.filter_by(username='admin').first():
        admin = User(username='admin', nombre='Administrador',
                     password=generate_password_hash('admin123'), rol='admin')
        db.session.add(admin)
        db.session.commit()

def login_required(f):
    from functools import wraps
    @wraps(f)
    def wrap(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return wrap

def admin_required(f):
    from functools import wraps
    @wraps(f)
    def wrap(*args, **kwargs):
        user = User.query.get(session['user_id'])
        if user.rol != 'admin':
            flash("Solo el admin tiene acceso a esta función.")
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return wrap

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username).first()
        if user and check_password_hash(user.password, password):
            session['user_id'] = user.id
            if user.rol == 'admin':
                return redirect(url_for('admin_panel'))
            else:
                return redirect(url_for('index'))
        flash('Usuario o contraseña incorrectos')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/register_user', methods=['GET', 'POST'])
@login_required
@admin_required
def register_user():
    if request.method == 'POST':
        nombre = request.form['nombre']
        username = request.form['username']
        password = request.form['password']
        rol = request.form['rol']
        if User.query.filter_by(username=username).first():
            flash('Nombre de usuario ya existe')
            return render_template('register_user.html')
        user = User(
            nombre=nombre,
            username=username,
            password=generate_password_hash(password),
            rol=rol
        )
        db.session.add(user)
        db.session.commit()
        flash('Usuario creado exitosamente')
        return redirect(url_for('admin_panel'))
    return render_template('register_user.html')

@app.route('/', methods=['GET', 'POST'])
@login_required
def index():
    user = User.query.get(session['user_id'])
    if user.rol == 'admin':
        return redirect(url_for('admin_panel'))

    if request.method == 'POST':
        nombre_evento = request.form['nombre_evento']
        telefono = request.form['telefono']
        lugar = request.form['lugar']
        fecha = request.form['fecha']
        fotos = request.files.getlist('fotos')

        if not (nombre_evento and telefono and lugar and fecha and fotos):
            flash("Completa todos los campos y agrega al menos una foto.")
            return render_template('index.html', user=user)
        img_paths = []
        for foto in fotos:
            if foto.filename == '':
                continue
            ext = foto.filename.split('.')[-1]
            filename = secrets.token_hex(8) + "." + ext
            file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            foto.save(file_path)
            img_paths.append(file_path)
        fotos_str = ';'.join(img_paths)

        informe = Informe(
            nombre_evento=nombre_evento,
            telefono=telefono,
            lugar=lugar,
            fecha=fecha,
            fotos=fotos_str,
            autor=user
        )
        db.session.add(informe)
        db.session.commit()
        flash('Informe guardado correctamente')
        return redirect(url_for('trabajador_historial'))
    return render_template('index.html', user=user)

@app.route('/trabajador_historial')
@login_required
def trabajador_historial():
    user = User.query.get(session['user_id'])
    informes = Informe.query.filter_by(user_id=user.id).order_by(Informe.creado.desc()).all()
    return render_template('trabajador_historial.html', informes=informes, user=user)

@app.route('/descargar_word/<int:informe_id>')
@login_required
def descargar_word(informe_id):
    informe = Informe.query.get_or_404(informe_id)
    if informe.user_id != session['user_id'] and User.query.get(session['user_id']).rol != "admin":
        flash("No tienes permisos para descargar este informe.")
        return redirect(url_for('index'))

    doc = Document()
    doc.add_heading('Informe de Evento', 0)
    doc.add_paragraph(f'Nombre del evento: {informe.nombre_evento}')
    doc.add_paragraph(f'Nombre encargado: {informe.autor.nombre}')
    doc.add_paragraph(f'Teléfono: {informe.telefono}')
    doc.add_paragraph(f'Lugar: {informe.lugar}')
    doc.add_paragraph(f'Fecha: {informe.fecha}')
    doc.add_heading('Pruebas Fotográficas', level=1)

    fotos = [f for f in informe.fotos.split(';') if f]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    fotos_iter = iter(fotos)
    while True:
        try:
            row_cells = table.add_row().cells
            img1 = next(fotos_iter)
            run1 = row_cells[0].paragraphs[0].add_run()
            run1.add_picture(img1, width=Inches(3))
        except StopIteration:
            break
        try:
            img2 = next(fotos_iter)
            run2 = row_cells[1].paragraphs[0].add_run()
            run2.add_picture(img2, width=Inches(3))
        except StopIteration:
            pass

    nombre_doc = f"informe_{informe.id}.docx"
    doc_path = os.path.join(app.config["UPLOAD_FOLDER"], nombre_doc)
    doc.save(doc_path)
    return send_file(doc_path, as_attachment=True)

# Panel administrativo: visualizar informes y crear usuarios
@app.route('/admin_panel')
@login_required
@admin_required
def admin_panel():
    informes = Informe.query.order_by(Informe.creado.desc()).all()
    usuarios = User.query.filter(User.rol != 'admin').all()
    return render_template('admin_panel.html', informes=informes, usuarios=usuarios, user=User.query.get(session["user_id"]))

# Marcar un informe como pagado
@app.route('/marcar_pagado/<int:informe_id>')
@login_required
@admin_required
def marcar_pagado(informe_id):
    informe = Informe.query.get_or_404(informe_id)
    informe.pagado = True
    db.session.commit()
    flash("Informe marcado como pagado/comisionado")
    return redirect(url_for('admin_panel'))

# Eliminar informe
@app.route('/eliminar_informe/<int:informe_id>')
@login_required
@admin_required
def eliminar_informe(informe_id):
    informe = Informe.query.get_or_404(informe_id)
    # Elimina las fotos asociadas del sistema de archivos
    fotos = [f for f in informe.fotos.split(';') if f]
    for foto_path in fotos:
        if os.path.exists(foto_path):
            os.remove(foto_path)
    db.session.delete(informe)
    db.session.commit()
    flash("Informe eliminado correctamente.")
    return redirect(url_for('admin_panel'))
@app.route('/cambiar_contrasena_admin/<int:user_id>', methods=['GET', 'POST'])
@login_required
@admin_required
def cambiar_contrasena_admin(user_id):
    usuario = User.query.get_or_404(user_id)
    if request.method == 'POST':
        nueva = request.form['nueva']
        repetir = request.form['repetir']
        if nueva != repetir:
            flash('Las contraseñas nuevas no coinciden')
        elif len(nueva) < 4:
            flash('La nueva contraseña debe tener al menos 4 caracteres')
        else:
            usuario.password = generate_password_hash(nueva)
            db.session.commit()
            flash(f'Contraseña cambiada con éxito para {usuario.nombre}')
            return redirect(url_for('admin_panel'))
    return render_template('cambiar_contrasena_admin.html', usuario=usuario, user=User.query.get(session['user_id']))

if __name__ == '__main__':
    app.run(debug=True)